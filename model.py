import os
import logging
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
from datetime import datetime
from fuzzywuzzy import process
from googletrans import Translator 

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

API_KEY = os.getenv("GENAI_API_KEY", "AIzaSyASH5t4dKbMMwbXk2M6ARMQ5UVjnR_5ysE")
if not API_KEY:
    logger.error("Google API key is missing. Set the GENAI_API_KEY environment variable.")
    raise ValueError("Google API key not set.")

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel('gemini-1.5-flash')

translator = Translator()

SYMPTOMS_KEYWORDS = [
    "fever", "cough", "headache", "nausea", "vomiting", "pain", "allergy",
    "asthma", "diabetes", "hypertension", "fatigue", "sore throat", "chronic", 
    "acute", "medication", "treatment", "mental"
]

def is_valid_symptom(symptom: str) -> bool:
    symptom = symptom.lower()
    matches = process.extract(symptom, SYMPTOMS_KEYWORDS, limit=10)
    return matches and matches[0][1] >= 80

def generate_long_text(prompt: str) -> str:
    try:
        logger.info(f"Generating content for prompt: {prompt}")
        response = model.generate_content(prompt)
        return response.text.strip() if response else 'No response from AI model.'
    except Exception as e:
        logger.error(f"Error generating content for prompt '{prompt}': {e}")
        return 'An error occurred while generating content.'

def process_symptoms(symptoms: str) -> dict:
    prompts = {
        "Symptoms": f"Based on the following symptoms: {symptoms}, what advice can you give? Give only 100 words.",
        "Treatment": f"Provide treatment options for someone suffering from: {symptoms}. Be direct. Give only 100 words.",
        "Avoidance": f"Enumerate precautions for someone experiencing: {symptoms}. Give only 100 words.",
        "Lifestyle Modifications": f"What lifestyle modifications should someone consider for: {symptoms}? Give only 100 words.",
        "Monitoring Strategies": f"Suggest monitoring strategies for: {symptoms}. Give only 100 words."
    }
    return {key: generate_long_text(prompt) for key, prompt in prompts.items()}

def translate_text(text: str, dest_lang: str) -> str:
    try:
        translated = translator.translate(text, dest=dest_lang)
        return translated.text
    except Exception as e:
        logger.error(f"Error translating text '{text}': {e}")
        return text 

@app.route('/healthcheck', methods=['POST'])
def healthcheck():
    try:
        data = request.json
        health_conditions = data.get('health_conditions', '').strip()
        target_language = data.get('target_language', 'en').strip()
        logger.info(f"Received health conditions: {health_conditions} and target language: {target_language}")

        if not health_conditions:
            return jsonify({"error": "No health conditions provided."}), 400

        health_conditions_en = translate_text(health_conditions, 'en')
        logger.info(f"Translated health conditions to English: {health_conditions_en}")

        if not is_valid_symptom(health_conditions_en):
            return jsonify({"error": "Please provide valid medical conditions only."}), 400

        advice_sections = process_symptoms(health_conditions_en)
        logger.info(f"Generated advice sections for: {health_conditions_en}")

        translated_advice = {key: translate_text(value, target_language) for key, value in advice_sections.items()}
        
        return jsonify({
            "title": "Health Advice",
            "symptoms": translated_advice["Symptoms"],
            "treatment": translated_advice["Treatment"],
            "avoidance": translated_advice["Avoidance"],
            "lifestyle_modifications": translated_advice["Lifestyle Modifications"],
            "monitoring_strategies": translated_advice["Monitoring Strategies"]
        })

    except Exception as e:
        logger.error(f"Error processing healthcheck request: {e}")
        return jsonify({"error": "An error occurred while processing the request."}), 500

@app.route('/download-advice', methods=['POST'])
def download_advice():
    try:
        data = request.json
        health_conditions = data.get('health_conditions', '').strip()
        age = data.get('age', 'N/A')
        gender = data.get('gender', 'N/A')
        doctor_name = data.get('doctor_name', 'Dr. Lumina')
        qualification = data.get('qualification', 'Qualification')
        name = data.get('name', 'Not provided')
        target_language = data.get('target_language', 'en').strip()
        current_time = datetime.now()

        logger.info(f"Received download request for health conditions: {health_conditions}")

        if not health_conditions:
            return jsonify({"error": "No health conditions provided."}), 400

        health_conditions_en = translate_text(health_conditions, 'en')
        logger.info(f"Translated health conditions to English: {health_conditions_en}")

        if not is_valid_symptom(health_conditions_en):
            return jsonify({"error": "Please provide valid medical conditions only."}), 400

        advice_sections = process_symptoms(health_conditions_en)
        translated_advice_sections = {key: translate_text(value, target_language) for key, value in advice_sections.items()}
        
        # Translate the health conditions to the target language
        translated_health_conditions = translate_text(health_conditions, target_language)

        doc_stream = create_health_advice_doc(
            doctor_name, qualification, age, gender, name, current_time, translated_health_conditions, translated_advice_sections
        )

        return send_file(
            doc_stream,
            as_attachment=True,
            download_name='health_advice.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logger.error(f"Error generating health advice document: {e}")
        return jsonify({"error": "An error occurred while generating the document."}), 500

def create_health_advice_doc(doctor_name, qualification, age, gender, name, current_time, health_conditions, advice_sections):
    try:
        doc = Document()
        
        title = doc.add_heading(doctor_name, level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].font.size = Pt(18)
        title.runs[0].font.color.rgb = RGBColor(0, 102, 204)
        
        qualification_paragraph = doc.add_paragraph(qualification)
        qualification_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        qualification_paragraph.runs[0].font.size = Pt(12)
        qualification_paragraph.runs[0].font.italic = True
        qualification_paragraph.runs[0].font.color.rgb = RGBColor(80, 80, 80)

        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = f"Age: {age}"
        table.cell(0, 1).text = f"Gender: {gender}"
        table.cell(1, 0).text = f"Date: {current_time.strftime('%B %d, %Y')}"
        table.cell(1, 1).text = f"Time: {current_time.strftime('%I:%M %p')}"
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(60, 60, 60)

        doc.add_paragraph()

        sections = {
            "Patient Name": name,
            "Health Conditions": health_conditions,
            "Symptoms": advice_sections["Symptoms"],
            "Treatment": advice_sections["Treatment"],
            "Avoidance": advice_sections["Avoidance"],
            "Lifestyle Modifications": advice_sections["Lifestyle Modifications"],
            "Monitoring Strategies": advice_sections["Monitoring Strategies"]
        }

        for heading, content in sections.items():
            heading_paragraph = doc.add_paragraph(heading)
            heading_paragraph.runs[0].bold = True
            heading_paragraph.runs[0].font.size = Pt(12)
            heading_paragraph.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            content_paragraph = doc.add_paragraph(content)
            content_paragraph.paragraph_format.space_after = Pt(10)
            content_paragraph.runs[0].font.size = Pt(11)
            content_paragraph.runs[0].font.name = 'Arial'
        
        doc.add_paragraph()
        stamp_path = "stamp.png"
        if os.path.exists(stamp_path):
            doc.add_picture(stamp_path, width=docx.shared.Inches(2))
        
        doc.add_paragraph().add_run(f"Generated on {current_time.strftime('%B %d, %Y at %I:%M %p')}")

        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        return output_stream

    except Exception as e:
        logger.error(f"Error creating health advice document: {e}")
        return None

if __name__ == "__main__":
    app.run(host='0.0.0.0',port=5000,debug=True)
